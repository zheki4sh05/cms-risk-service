"""Generate Word document with functional capabilities and role access for cms-risk-service."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_data, header: bool = False) -> None:
    row = table.add_row()
    for idx, text in enumerate(cells_data):
        cell = row.cells[idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            if header:
                for run in paragraph.runs:
                    run.bold = True
        if header:
            set_cell_shading(cell, "D9E2F3")


def format_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.name = "Calibri"


def add_vertical_spec_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        label_cell, value_cell = row.cells[0], row.cells[1]
        format_cell_text(label_cell, label)
        format_cell_text(value_cell, value)
        set_cell_shading(label_cell, "D9E2F3")
        for run in label_cell.paragraphs[0].runs:
            run.bold = True


def add_compliance_rules_use_cases(doc: Document) -> None:
    doc.add_heading("Управлять правилами комплаенса — запросы и ответы", level=1)

    hdr_get = (
        "Заголовки: Authorization: Bearer <token> (обязательно); "
        "CompanyId: <uuid> (необязательно, если companyId есть в JWT)."
    )
    hdr_body = (
        "Заголовки: Authorization: Bearer <token>; CompanyId: <uuid> (или из JWT); "
        "Content-Type: application/json."
    )
    err_common = "401 — нет или неверный токен; 400 — не указан CompanyId; 500 — ошибка сервера."
    err_by_id = err_common + " 404 — правило не найдено в компании пользователя."
    err_write_create = err_common + " 403 — нет права MANAGE_RULES_AND_RISKS. 400 — ошибка валидации или неверный JSON."
    err_write_update = err_by_id + " 403 — нет права MANAGE_RULES_AND_RISKS. 400 — ошибка валидации тела."

    cases = [
        (
            "Список правил",
            "GET /api/rules",
            f"{hdr_get}\n"
            "Параметры: нет (ни в пути, ни в query, ни в теле).\n"
            "Выборка по companyId из контекста пользователя.",
            "200 — объект с массивом items: id, name, condition, action, categoryId, "
            "categoryLabel, priority, enabled, riskObjectId, riskObject.",
            err_common,
        ),
        (
            "Детали правила",
            "GET /api/rules/{id}",
            f"{hdr_get}\n"
            "Параметр пути: id — UUID правила.\n"
            "Тело: нет.",
            "200 — id, companyId, name, condition, categoryId, riskObjectId, priority, "
            "responsibleUserId, actions[], enabled, mechanismScriptName, "
            "mechanismScriptContent, createdByUserId, savedAt.",
            err_by_id,
        ),
        (
            "Создать правило",
            "POST /api/rules",
            f"{hdr_body}\n"
            "Тело JSON, обязательно: name, condition, categoryId, priority (low | medium | high), "
            "actions[] (минимум одно: createIncident или sendNotification), enabled.\n"
            "Необязательно: riskObjectId, responsibleUserId, mechanismScriptName, mechanismScriptContent.",
            "201 — id (новое правило), savedAt.",
            err_write_create,
        ),
        (
            "Изменить правило",
            "PUT /api/rules/{id}",
            f"{hdr_body}\n"
            "Параметр пути: id — UUID правила.\n"
            "Тело JSON, обязательно: description (текст для истории изменений), name, condition, "
            "categoryId, priority, actions[], enabled.\n"
            "Необязательно: riskObjectId, responsibleUserId, mechanismScriptName, mechanismScriptContent.",
            "200 — id, savedAt (предыдущая версия пишется в историю).",
            err_write_update,
        ),
        (
            "Изменить привязку объекта риска",
            "PUT /api/rules/{id}/risk-object",
            f"{hdr_body}\n"
            "Параметр пути: id — UUID правила.\n"
            "Тело JSON: одно поле riskObjectId (UUID или null — снять привязку). "
            "Остальные поля правила не передаются.",
            "200 — id правила, savedAt.",
            err_write_update,
        ),
    ]

    for i, (title, req, desc, resp, err) in enumerate(cases, 1):
        doc.add_heading(f"{i}. {title}", level=2)
        add_vertical_spec_table(
            doc,
            [
                ("Вариант использования", title),
                ("Запрос", req),
                ("Описание запроса", desc),
                ("Ответ", resp),
                ("Исключительные ситуации", err),
            ],
        )
        doc.add_paragraph()


def main() -> None:
    doc = Document()
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)
    doc.sections[0].left_margin = Cm(2.5)
    doc.sections[0].right_margin = Cm(2)

    title = doc.add_heading("Перечень функциональных возможностей", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("cms-risk-service")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Доступ: JWT + CompanyId. Изменение правил и категорий — только с правом MANAGE_RULES_AND_RISKS."
    )

    doc.add_heading("Роли", level=1)
    roles = doc.add_table(rows=1, cols=3)
    roles.style = "Table Grid"
    rh = roles.rows[0].cells
    rh[0].text, rh[1].text, rh[2].text = "Роль", "Название", "Доступ"
    for c in rh:
        set_cell_shading(c, "D9E2F3")
    for row in [
        ("EXECUTIVE", "Руководитель", "Все права, в т.ч. управление правилами"),
        ("SUPERVISOR", "Руководитель подразделения", "По назначенным правам"),
        ("MANAGER", "Менеджер", "По назначенным правам"),
        ("DEFAULT", "Пользователь", "По назначенным правам"),
    ]:
        add_table_row(roles, row)

    doc.add_heading("Функции сервиса", level=1)
    cap = doc.add_table(rows=1, cols=3)
    cap.style = "Table Grid"
    ch = cap.rows[0].cells
    ch[0].text, ch[1].text, ch[2].text = "Функция", "API", "Кто может"
    for c in ch:
        set_cell_shading(c, "D9E2F3")
    for row in [
        ("Список и просмотр правил", "GET /api/rules, /api/rules/{id}", "Любой с JWT"),
        ("Создание и изменение правил", "POST/PUT /api/rules", "MANAGE_RULES_AND_RISKS"),
        ("Привязка объекта риска", "PUT /api/rules/{id}/risk-object", "MANAGE_RULES_AND_RISKS"),
        ("История изменений", "GET /api/rules/change-history", "Любой с JWT"),
        ("Категории рисков", "GET/POST/PUT/DELETE /api/risk-categories", "Чтение — JWT; запись — MANAGE_RULES_AND_RISKS"),
        ("Статистика обработки", "GET /api/rules/processing/statistic", "Любой с JWT"),
        ("Результаты верификации", "GET/POST/PUT/DELETE /api/verification-results", "Любой с JWT"),
        ("Фоновая обработка", "Kafka, планировщики", "Система, без пользователя"),
    ]:
        add_table_row(cap, row)

    doc.add_heading("Права (permissions)", level=1)
    perm = doc.add_table(rows=1, cols=2)
    perm.style = "Table Grid"
    ph = perm.rows[0].cells
    ph[0].text, ph[1].text = "Право", "Что даёт"
    for c in ph:
        set_cell_shading(c, "D9E2F3")
    for row in [
        ("VIEW_RULES_AND_RISKS_PAGE", "Раздел «Правила и риски» в UI"),
        ("MANAGE_RULES_AND_RISKS", "Создание и изменение правил и категорий (проверяет cms-risk-service)"),
        ("VIEW_RISK_OBJECTS_PAGE", "Объекты риска в UI (другой сервис)"),
        ("MANAGE_RISK_OBJECTS", "Редактирование объектов риска (другой сервис)"),
        ("EDIT_USERS", "Назначение прав другим пользователям"),
        ("VIEW_*_PAGE (остальные)", "Доступ к разделам UI: дашборд, пользователи, интеграции, настройки, профиль"),
        ("MANAGE_INTEGRATIONS", "Настройка интеграций"),
    ]:
        add_table_row(perm, row)

    doc.add_paragraph("EXECUTIVE при регистрации получает все права. Остальным права выдаёт администратор (EDIT_USERS).")

    add_compliance_rules_use_cases(doc)

    foot = doc.add_paragraph("cms-risk-service, 17.05.2026")
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].italic = True

    import os

    output_path = (
        r"C:\trustflow-backend\cms-risk-service\docs"
        r"\Перечень_функциональных_возможностей_cms-risk-service.docx"
    )
    # Если основной файл открыт в Word — сохраняем копию
    fallback_path = (
        r"C:\trustflow-backend\cms-risk-service\docs"
        r"\Перечень_функциональных_возможностей_cms-risk-service_краткий.docx"
    )
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    try:
        doc.save(output_path)
        print(output_path)
    except PermissionError:
        doc.save(fallback_path)
        print(fallback_path)
        print("(основной файл занят — закройте Word и перезапустите скрипт)")


if __name__ == "__main__":
    main()
